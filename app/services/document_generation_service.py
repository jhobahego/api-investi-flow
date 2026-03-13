"""
Servicio para generar documentos .docx a partir de HTML
Compatible con TipTap editor output
"""

import importlib.util
import os
from typing import List

if importlib.util.find_spec("docx") is None or importlib.util.find_spec("lxml") is None:
    raise ImportError(
        "Las dependencias python-docx o lxml no están instaladas. "
        "Ejecuta: pip install python-docx lxml"
    )

import lxml.html
from docx import Document
from fastapi import HTTPException


class DocumentGenerationService:
    """Servicio para generar un archivo .docx a partir de HTML"""

    @staticmethod
    def html_pages_to_docx(pages: List[str], output_path: str) -> None:
        """
        Toma una lista de páginas HTML y genera o sobreescribe un archivo .docx.

        Args:
            pages: Lista de strings HTML
            output_path: Ruta de destino para guardar el .docx
        """
        try:
            document = Document()

            # Unir todas las páginas en un string HTML
            full_html = "".join(pages)
            if not full_html.strip():
                document.save(output_path)
                return

            fragments = lxml.html.fragments_fromstring(full_html)

            for element in fragments:
                if isinstance(element, str):
                    if element.strip():
                        p = document.add_paragraph()
                        p.add_run(element.strip())
                    continue

                DocumentGenerationService._process_element(
                    document, element, list_level=0
                )

            # Ensure parent directories exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            document.save(output_path)

        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error al generar el documento .docx: {str(e)}"
            )

    @staticmethod
    def _process_element(document_or_parent, element, list_level=0, list_style=None):
        tag = element.tag.lower()

        # Heading 1
        if tag == "h1":
            p = document_or_parent.add_paragraph(style="Heading 1")
            DocumentGenerationService._process_inline_elements(p, element)
        # Heading 2
        elif tag == "h2":
            p = document_or_parent.add_paragraph(style="Heading 2")
            DocumentGenerationService._process_inline_elements(p, element)
        # Heading 3
        elif tag == "h3" or tag == "h4" or tag == "h5" or tag == "h6":
            p = document_or_parent.add_paragraph(style="Heading 3")
            DocumentGenerationService._process_inline_elements(p, element)
        # Paragraph
        elif tag == "p":
            text_content = element.text_content().strip()
            if text_content or len(element) > 0:
                p = document_or_parent.add_paragraph()
                DocumentGenerationService._process_inline_elements(p, element)
            else:
                # Empty paragraph
                document_or_parent.add_paragraph()
        # Lists
        elif tag in ["ul", "ol"]:
            style = "List Bullet" if tag == "ul" else "List Number"
            for child in element:
                if child.tag.lower() == "li":
                    p = document_or_parent.add_paragraph(style=style)
                    # We could adjust left indent based on list_level if needed
                    # but python-docx styles handle standard levels
                    DocumentGenerationService._process_inline_elements(p, child)

                    # Process nested lists inside li
                    for nested in child:
                        if nested.tag.lower() in ["ul", "ol"]:
                            DocumentGenerationService._process_element(
                                document_or_parent, nested, list_level + 1
                            )
        # Table (Minimal Support)
        elif tag == "table":
            rows = element.findall(".//tr")
            if not rows:
                return

            # Find max cols
            max_cols = max(
                [len(tr.findall(".//td") + tr.findall(".//th")) for tr in rows],
                default=1,
            )
            table = document_or_parent.add_table(rows=0, cols=max_cols)
            table.style = "Table Grid"

            for tr in rows:
                cells = tr.findall(".//td") + tr.findall(".//th")
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < max_cols:
                        DocumentGenerationService._process_inline_elements(
                            row_cells[i].paragraphs[0], cell
                        )

        else:
            # Fallback for unrecognized tags like divs or spans at root level
            if element.text_content().strip():
                p = document_or_parent.add_paragraph()
                DocumentGenerationService._process_inline_elements(p, element)

    @staticmethod
    def _process_inline_elements(paragraph, element):
        """Processes elements that live inside a paragraph (text, bold, italic, etc)"""
        if element.text and element.text.strip():
            paragraph.add_run(element.text)

        for child in element:
            tag = child.tag.lower()
            run = paragraph.add_run()

            if tag in ["strong", "b"]:
                run.bold = True
            elif tag in ["em", "i"]:
                run.italic = True
            elif tag in ["u"]:
                run.underline = True

            if child.text:
                run.text = child.text

                # Check formatting of the run
            if child.tail:
                paragraph.add_run(child.tail)


document_generation_service = DocumentGenerationService()
