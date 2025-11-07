"""
Servicio para extracción de contenido de documentos .docx a HTML
Compatible con TipTap editor
"""
import re
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError(
        "python-docx no está instalado. " "Ejecuta: pip install python-docx"
    )

from fastapi import HTTPException


class DocumentExtractionService:
    """Servicio para extraer contenido de documentos y convertirlo a HTML"""

    # Configuración de páginas
    CHARS_PER_PAGE = 3000  # Aproximadamente 3000 caracteres por página A4

    @staticmethod
    def extract_docx_to_html(file_path: str) -> str:
        """
        Extrae el contenido de un archivo .docx y lo convierte a HTML
        compatible con TipTap editor

        Args:
            file_path: Ruta al archivo .docx

        Returns:
            str: Contenido HTML del documento

        Raises:
            HTTPException: Si el archivo no existe o no se puede leer
        """
        path = Path(file_path)

        if not path.exists():
            raise HTTPException(
                status_code=404, detail=f"Archivo no encontrado: {file_path}"
            )

        if not path.suffix.lower() == ".docx":
            raise HTTPException(
                status_code=400, detail="Solo se soportan archivos .docx"
            )

        try:
            document = Document(file_path)
            html_content = DocumentExtractionService._convert_document_to_html(document)
            return html_content

        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error al extraer contenido del documento: {str(e)}",
            )

    @staticmethod
    def extract_docx_to_pages(file_path: str) -> List[str]:
        """
        Extrae el contenido de un archivo .docx y lo divide en páginas HTML

        Args:
            file_path: Ruta al archivo .docx

        Returns:
            List[str]: Lista de páginas HTML del documento

        Raises:
            HTTPException: Si el archivo no existe o no se puede leer
        """
        path = Path(file_path)

        if not path.exists():
            raise HTTPException(
                status_code=404, detail=f"Archivo no encontrado: {file_path}"
            )

        if not path.suffix.lower() == ".docx":
            raise HTTPException(
                status_code=400, detail="Solo se soportan archivos .docx"
            )

        try:
            document = Document(file_path)
            pages = DocumentExtractionService._convert_document_to_pages(document)
            return pages

        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error al extraer contenido del documento: {str(e)}",
            )

    @staticmethod
    def _convert_document_to_html(document: Document) -> str:
        """
        Convierte un objeto Document de python-docx a HTML

        Args:
            document: Objeto Document de python-docx

        Returns:
            str: HTML generado
        """
        html_parts = []

        for element in document.element.body:
            # Procesar párrafos
            if element.tag.endswith("p"):
                para = Paragraph(element, document)
                html_parts.append(
                    DocumentExtractionService._paragraph_to_html(para, document)
                )
            # Procesar tablas
            elif element.tag.endswith("tbl"):
                table = Table(element, document)
                html_parts.append(DocumentExtractionService._table_to_html(table))

        # Si no hay contenido, devolver párrafo vacío
        if not html_parts or all(part.strip() == "" for part in html_parts):
            return "<p></p>"

        # Agrupar elementos de lista consecutivos
        grouped_html = DocumentExtractionService._group_list_items(html_parts)

        return "".join(grouped_html)

    @staticmethod
    def _convert_document_to_pages(document: Document) -> List[str]:
        """
        Convierte un objeto Document de python-docx a lista de páginas HTML

        Args:
            document: Objeto Document de python-docx

        Returns:
            List[str]: Lista de páginas HTML
        """
        all_elements = []

        # Recolectar todos los elementos HTML
        for element in document.element.body:
            # Procesar párrafos
            if element.tag.endswith("p"):
                para = Paragraph(element, document)
                html = DocumentExtractionService._paragraph_to_html(para, document)
                if html.strip():
                    all_elements.append(html)
            # Procesar tablas
            elif element.tag.endswith("tbl"):
                table = Table(element, document)
                html = DocumentExtractionService._table_to_html(table)
                if html.strip():
                    all_elements.append(html)

        # Si no hay contenido, devolver una página vacía
        if not all_elements:
            return ["<p></p>"]

        # Agrupar elementos de lista consecutivos
        grouped_elements = DocumentExtractionService._group_list_items(all_elements)

        # Dividir en páginas
        pages: List[str] = []
        current_page: List[str] = []
        current_length = 0

        for element_html in grouped_elements:
            # Calcular longitud de texto del elemento
            element_text = re.sub(r"<[^>]+>", "", element_html)
            element_length = len(element_text)

            # Si agregar este elemento excede el límite, crear nueva página
            if (
                current_length + element_length
                > DocumentExtractionService.CHARS_PER_PAGE
                and current_page
            ):
                pages.append("".join(current_page))
                current_page = [element_html]
                current_length = element_length
            else:
                current_page.append(element_html)
                current_length += element_length

        # Agregar última página
        if current_page:
            pages.append("".join(current_page))

        return pages if pages else ["<p></p>"]

    @staticmethod
    def _paragraph_to_html(paragraph: Paragraph, document: Document) -> str:
        """
        Convierte un párrafo de docx a HTML con formato

        Args:
            paragraph: Objeto Paragraph de python-docx
            document: Objeto Document de python-docx (para acceder al numbering)

        Returns:
            str: HTML del párrafo con metadata de tipo de lista si aplica
        """
        text = paragraph.text.strip()

        # Si el párrafo está vacío, devolver un párrafo HTML vacío
        if not text:
            return "<p></p>"

        # Detectar nivel de encabezado
        style = paragraph.style.name.lower() if paragraph.style else ""

        if "heading 1" in style or "título 1" in style:
            return f"<h1>{DocumentExtractionService._format_runs(paragraph)}</h1>"
        elif "heading 2" in style or "título 2" in style:
            return f"<h2>{DocumentExtractionService._format_runs(paragraph)}</h2>"
        elif "heading 3" in style or "título 3" in style:
            return f"<h3>{DocumentExtractionService._format_runs(paragraph)}</h3>"

        # Detectar listas
        if paragraph._element.pPr is not None:
            numPr = paragraph._element.pPr.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if numPr is not None:
                # Es un elemento de lista
                ilvl = numPr.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl"
                )
                numId = numPr.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId"
                )

                if ilvl is not None and numId is not None:
                    # Obtener el valor del numId
                    num_id_val = numId.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    )
                    ilvl_val = ilvl.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    )
                    # Determinar si es lista ordenada o no ordenada
                    list_type = DocumentExtractionService._detect_list_type(
                        document, num_id_val, ilvl_val
                    )
                    # Usar atributo data-list-type para identificar el tipo de lista
                    return f"<li data-list-type='{list_type}' data-num-id='{num_id_val}'>{DocumentExtractionService._format_runs(paragraph)}</li>"

        # Párrafo normal
        return f"<p>{DocumentExtractionService._format_runs(paragraph)}</p>"

    @staticmethod
    def _format_runs(paragraph: Paragraph) -> str:
        """
        Procesa los runs de un párrafo aplicando formato en línea

        Args:
            paragraph: Objeto Paragraph de python-docx

        Returns:
            str: Texto con formato HTML
        """
        html_text = []

        for run in paragraph.runs:
            text = run.text

            # Aplicar formato en línea
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"

            html_text.append(text)

        return "".join(html_text)

    @staticmethod
    def _table_to_html(table: Table) -> str:
        """
        Convierte una tabla de docx a HTML

        Args:
            table: Objeto Table de python-docx

        Returns:
            str: HTML de la tabla
        """
        html = ["<table>"]

        for i, row in enumerate(table.rows):
            html.append("<tr>")
            for cell in row.cells:
                tag = "th" if i == 0 else "td"
                cell_text = cell.text.strip()
                html.append(f"<{tag}>{cell_text}</{tag}>")
            html.append("</tr>")

        html.append("</table>")
        return "".join(html)

    @staticmethod
    def _detect_list_type(
        document: Document, num_id: Optional[str], ilvl: Optional[str]
    ) -> str:
        """
        Detecta si una lista es ordenada (ol) o no ordenada (ul) accediendo al numbering.xml

        Args:
            document: Objeto Document de python-docx
            num_id: ID de numeración del párrafo
            ilvl: Nivel de indentación

        Returns:
            str: 'ol' para lista ordenada, 'ul' para lista no ordenada
        """
        # Intentar acceder al numbering part para obtener el numFmt real
        try:
            if (
                hasattr(document.part, "numbering_part")
                and document.part.numbering_part
            ):
                numbering_part = document.part.numbering_part
                numbering_element = numbering_part._element
                ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

                if num_id:
                    # Buscar el elemento num con el numId
                    num = numbering_element.find(f'.//{ns}num[@{ns}numId="{num_id}"]')

                    if num is not None:
                        # Obtener abstractNumId
                        abstract_num_id_el = num.find(f".//{ns}abstractNumId")
                        if abstract_num_id_el is not None:
                            abstract_num_id = abstract_num_id_el.get(f"{ns}val")

                            # Buscar abstractNum
                            abstract_num = numbering_element.find(
                                f'.//{ns}abstractNum[@{ns}abstractNumId="{abstract_num_id}"]'
                            )

                            if abstract_num is not None:
                                # Buscar el nivel correcto (por defecto nivel 0)
                                lvl_val = ilvl if ilvl else "0"
                                lvl_el = abstract_num.find(
                                    f'.//{ns}lvl[@{ns}ilvl="{lvl_val}"]'
                                )

                                if lvl_el is not None:
                                    # Buscar numFmt
                                    num_fmt = lvl_el.find(f".//{ns}numFmt")
                                    if num_fmt is not None:
                                        fmt_val = num_fmt.get(f"{ns}val")

                                        # Formatos numéricos = lista ordenada
                                        if fmt_val in [
                                            "decimal",
                                            "upperRoman",
                                            "lowerRoman",
                                            "upperLetter",
                                            "lowerLetter",
                                            "decimalZero",
                                            "ordinal",
                                        ]:
                                            return "ol"
                                        # Formatos de viñetas = lista no ordenada
                                        elif fmt_val == "bullet":
                                            return "ul"
        except Exception as e:
            # Log error but continue with fallback
            import logging

            logging.warning(f"Error detecting list type from numbering.xml: {e}")

        # Fallback: Por defecto asumir lista con viñetas (más seguro que asumir numerada)
        return "ul"

    @staticmethod
    def _group_list_items(html_elements: List[str]) -> List[str]:
        """
        Agrupa elementos <li> consecutivos dentro de etiquetas <ul> o <ol>

        Args:
            html_elements: Lista de elementos HTML

        Returns:
            List[str]: Lista de elementos HTML con listas agrupadas
        """
        grouped = []
        current_list_items = []
        current_list_type = None

        for element in html_elements:
            # Verificar si es un elemento de lista
            if element.startswith("<li"):
                # Extraer el tipo de lista del atributo data-list-type
                match = re.search(r"data-list-type='(\w+)'", element)
                list_type = match.group(1) if match else "ul"

                # Si cambia el tipo de lista, cerrar la lista anterior
                if current_list_type is not None and current_list_type != list_type:
                    # Cerrar lista anterior
                    grouped.append(f"<{current_list_type}>")
                    for item in current_list_items:
                        # Remover el atributo data-list-type antes de agregar
                        clean_item = re.sub(r"\s*data-list-type='[^']*'", "", item)
                        grouped.append(clean_item)
                    grouped.append(f"</{current_list_type}>")
                    current_list_items = []

                # Agregar item a la lista actual
                current_list_items.append(element)
                current_list_type = list_type
            else:
                # No es un elemento de lista
                # Si había elementos de lista acumulados, cerrar la lista
                if current_list_items:
                    grouped.append(f"<{current_list_type}>")
                    for item in current_list_items:
                        # Remover el atributo data-list-type antes de agregar
                        clean_item = re.sub(r"\s*data-list-type='[^']*'", "", item)
                        grouped.append(clean_item)
                    grouped.append(f"</{current_list_type}>")
                    current_list_items = []
                    current_list_type = None

                # Agregar el elemento normal
                grouped.append(element)

        # Si quedan elementos de lista al final, cerrar la lista
        if current_list_items:
            grouped.append(f"<{current_list_type}>")
            for item in current_list_items:
                # Remover el atributo data-list-type antes de agregar
                clean_item = re.sub(r"\s*data-list-type='[^']*'", "", item)
                grouped.append(clean_item)
            grouped.append(f"</{current_list_type}>")

        return grouped

    @staticmethod
    def get_document_preview(file_path: str, max_chars: int = 200) -> str:
        """
        Obtiene una vista previa del contenido del documento (texto plano)

        Args:
            file_path: Ruta al archivo .docx
            max_chars: Número máximo de caracteres para la vista previa

        Returns:
            str: Vista previa del documento
        """
        try:
            document = Document(file_path)
            full_text = "\n".join([para.text for para in document.paragraphs])

            if len(full_text) <= max_chars:
                return full_text

            return full_text[:max_chars] + "..."

        except Exception as e:
            return f"Error al obtener vista previa: {str(e)}"


# Instancia del servicio
document_extraction_service = DocumentExtractionService()
